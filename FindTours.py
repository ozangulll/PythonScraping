import os
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document

url = "https://www.excursionmarket.com/all-tour-fields" 
response = requests.get(url)
html_content = response.content
soup = BeautifulSoup(html_content, "html.parser")

cities = [city.text.strip() for city in soup.find_all("h4", class_="text-dark-1 text-18 fw-500")]

for city in cities:
    url = f"https://www.excursionmarket.com/search?q={city}"
    response = requests.get(url)
    html_content = response.content
    soup = BeautifulSoup(html_content, "html.parser")

    names = soup.find_all("h4", class_="tourCard__title text-dark-1 text-18 lh-16 fw-500")
    prices = soup.find_all("span", class_="text-16 fw-500 text-dark-1")
    times = soup.find_all("div", class_="d-flex items-center lh-14 mb-5")

    document = Document()
    document.add_heading(f"Tours in {city}", level=1)

    for i in range(len(names)):
        name = names[i].text.strip()
        price = prices[i].text.strip()
        time = times[i].text.strip().split('\n')[0]
        document.add_paragraph(f"Tour Name: {name}")
        document.add_paragraph(f"Price: {price}")
        document.add_paragraph(f"Times: {time}")
        document.add_paragraph("-------------------")

    desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    docx_path = os.path.join(desktop_path, "py", f"{city}.docx")

    document.save(docx_path)
    print(f"DOCX file saved to: {docx_path}")
