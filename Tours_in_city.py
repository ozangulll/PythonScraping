import os
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document


city = input("Please enter the city name: ")

url = f"https://www.excursionmarket.com/search?q={city}"
response = requests.get(url)
html_content = response.content
soup = BeautifulSoup(html_content, "html.parser")

names = soup.find_all("h4", attrs={"class": "tourCard__title text-dark-1 text-18 lh-16 fw-500"})
prices = soup.find_all("span", attrs={"class": "text-16 fw-500 text-dark-1"})
times_and_tourtypes = soup.find_all("div", attrs={"class": "d-flex items-center lh-14 mb-5"})

data = []

for i in range(len(names)):
    name = names[i].text.strip()
    price = prices[i].text.strip()
    times_and_tourtype = times_and_tourtypes[i].text.strip().split('\n')[0]  
    data.append([name, price, times_and_tourtype])

result = pd.DataFrame(data, columns=["Name of the tours", "Price", "Times and Tour Type"])
print(result)

desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
docx_path = os.path.join(desktop_path, "py", f"{city}.docx")


doc = Document()
doc.add_heading(f"Tours in {city}", level=1)

table = doc.add_table(result.shape[0]+1, result.shape[1])

hdr_cells = table.rows[0].cells
for i, column in enumerate(result.columns):
    hdr_cells[i].text = column

for i in range(result.shape[0]):
    for j in range(result.shape[-1]):
        table.cell(i+1, j).text = str(result.values[i, j])


doc.save(docx_path)
print(f"Document saved to: {docx_path}")
