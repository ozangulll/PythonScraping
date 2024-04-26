import os
import requests
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from docx import Document
import time


show_more_clicks = 0

driver = webdriver.Chrome()  # Chrome WebDriver'ı başlat
url = "https://www.excursionmarket.com/all-tours"
driver.get(url)


while show_more_clicks < 29:
    try:
        show_more_button = WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.ID, "tourLoadBtn")))
        driver.execute_script("arguments[0].click();", show_more_button)
        time.sleep(3)  
        show_more_clicks += 1  
    except:
        print("No more tours to load.")
        break


soup = BeautifulSoup(driver.page_source, "html.parser")


tours = [a['href'] for a in soup.select('.tourCard__image a[href]')]


for tour_url in tours:
    response = requests.get(tour_url)
    html_content = response.content
    soup = BeautifulSoup(html_content, "html.parser")

    faq_items = soup.find_all("div", class_="accordion__item")
    
    if not faq_items:
        print(f"No FAQ items found for {tour_url}")
        continue

    document = Document()
    document.add_heading(f"FAQs for {tour_url}", level=1)

    for faq_item in faq_items:
        question = faq_item.find("h3", class_="custom-h3-15")
        answer = faq_item.find("p", class_="text-15")

        if question and answer:
            question_text = question.text.strip()
            answer_text = answer.text.strip()
            document.add_paragraph(f"{question_text}")
            document.add_paragraph(f"{answer_text}")
            document.add_paragraph("-------------------")
        else:
            print("Incomplete FAQ item found.")

  
    tour_name = tour_url.split('/')[-1].replace('-', '_').replace(':', '')
    docx_file_path = os.path.join(os.environ['USERPROFILE'], 'Desktop', f"{tour_name}.docx")

  
    document.save(docx_file_path)
    print(f"DOCX file saved to: {docx_file_path}")


print(f"Number of clicks on 'show more' button: {show_more_clicks}")


driver.quit()
