import os
import requests
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from docx import Document
import time

show_more_clicks = 0

driver = webdriver.Chrome() 
url = "https://www.excursionmarket.com/all-tours"
driver.get(url)

while show_more_clicks < 29:
    try:
        show_more_button = WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.ID, "tourLoadBtn")))
        driver.execute_script("arguments[0].click();", show_more_button)
        time.sleep(3)  
        show_more_clicks += 1  
    except:
        print("No more tours to load.")
        break

soup = BeautifulSoup(driver.page_source, "html.parser")
tours = [a['href'] for a in soup.select('.tourCard__image a[href]')]


consolidated_document = Document()

for tour_url in tours:
    response = requests.get(tour_url)
    html_content = response.content
    soup = BeautifulSoup(html_content, "html.parser")

  
    consolidated_document.add_heading("Tour Link:", level=2)
    consolidated_document.add_paragraph(tour_url)

    includes = soup.find("div", class_="row x-gap-40 y-gap-40 pt-40")
    not_included = soup.find("div", class_="mt-40 border-top-light")

    if not includes or not not_included:
        print(f"Could not find 'Includes' or 'Not included' section for {tour_url}")
        continue

    includes_title = includes.find("h3").text.strip()
    not_included_title = not_included.find("h3").text.strip()

    includes_content = includes.find("div", class_="row x-gap-40 y-gap-40 pt-20")
    not_included_content = not_included.find("div", class_="row x-gap-40 y-gap-40 pt-20")

   
    consolidated_document.add_heading(includes_title, level=2)
    includes_items = includes_content.find_all("div", class_="col-12 text-dark-1 text-15 re-order")
    for item in includes_items:
        item_text = item.text.strip()
        consolidated_document.add_paragraph(item_text)

    bring_title = "What to bring"
    bring_content = soup.find("div", class_="mt-40 border-top-light").find_next_sibling("div")
    if bring_content:
        consolidated_document.add_heading(bring_title, level=2)
        bring_items = bring_content.find_all("div", class_="col-12 text-dark-1 text-15 re-order-neler-dahil-degil")
        for item in bring_items:
            item_text = item.text.strip()
            consolidated_document.add_paragraph(item_text)


    consolidated_document.add_heading(not_included_title, level=2)
    not_included_items = not_included_content.find_all("div", class_="col-12 text-dark-1 text-15 re-order-neler-dahil-degil")
    for item in not_included_items:
        item_text = item.text.strip()
        consolidated_document.add_paragraph(item_text)

print(f"Number of clicks on 'show more' button: {show_more_clicks}")

consolidated_docx_file_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop', "py", "consolidated_include_notinclude_bring.docx")
consolidated_document.save(consolidated_docx_file_path)
print(f"Consolidated DOCX file saved to: {consolidated_docx_file_path}")

driver.quit()
